import re
import os
import sys
import docx2txt
import docx
import json as js
import string
import random
letters = string.ascii_lowercase




def checkFontStyles(i):
    data_dic={}
    if(i.bold):
        data_dic['bold']=True

    if(i.highlight_color):
        data_dic['highlightColor']=True

    if(i.italic):
        data_dic['italic']=True

    if(i.outline):
        data_dic['outline']=True

    if(i.strike):
        data_dic['strike']=True

    if(i.underline):
        data_dic['underline']=True

    # new items 
    # 
    if(i.imprint):
        data_dic['imprint']=True
    # if(i.color):
    #     data_dic['color']=i.color.theme_color
    #     print(data_dic['color'])
    if(i.cs_bold):
        data_dic['cs_bold']=True
    if(i.cs_italic):
        data_dic['cs_italic']=True
    if(i.double_strike):
        data_dic['double_strike']=True
    if(i.emboss):
        data_dic['emboss']=True
    if(i.hidden):
        data_dic['hidden']=True
    if(i.math):
        data_dic['math']=True
    if(i.name):
        data_dic['fontName']=i.name
        # print(data_dic['name'])
    if(i.no_proof):
        data_dic['no_proof']=True
    if(i.rtl):
        data_dic['rtl']=True
    if(i.shadow):
        data_dic['shadow']=True
    if(i.small_caps):
        data_dic['small_caps']=True
    if(i.snap_to_grid):
        data_dic['snap_to_grid']=True
    if(i.spec_vanish):
        data_dic['spec_vanish']=True
    if(i.subscript):
        data_dic['subscript']=True
    if(i.superscript):
        data_dic['superscript']=True
    if(i.web_hidden):
        data_dic['web_hidden']=True
    # 
    #  End 

    if(len(data_dic) >=1):
        return data_dic
    else:
        data_dic['normal']=True
        return data_dic

    


#merge duplicate styles and compain text
def mergeTextbyStyles(ls):
    indexList=[]
    for index,obj in enumerate(ls):
        cur = obj['style'] #current
        # print(cur)
        if index > 0 :
            pre = ls[index -1]['style'] # previous
            if(cur == pre) :
                indexList.append(index-1)
                indexList.append(index)

    indexList=list(set(indexList))
    returnList=[]
    val=''
    for index,obj in enumerate(ls):
        dic={}
        if index in indexList:
            dic['style']=obj['style']
            val=val+obj['text']
            dic['text']=val
            if(index==indexList[-1]):
                returnList.append(dic)
        else:
            dic=obj
            returnList.append(dic.copy())   
    return returnList



#main converter

def word_to_json_parser(doc_path,store=False,image_folder="android",json="output.json"):
    ext=doc_path.split('.')
    if ext[-1].lower() != 'docx':
        print('[*] Warning! Unable to verify input file '+str(doc_path))
        exit()

    # print(ext)
    # break
    try:
        image_files=os.listdir(image_folder)
        if image_files!=[]:
            print('[*] Warning! Files alredy exist in '+str(image_folder)+' folder')
    except:
        os.mkdir(image_folder)
    docx2txt.process(doc_path, image_folder)
    image_files=os.listdir(image_folder)
    data_list=[]
    doc=docx.Document(doc_path)
    for val in doc.paragraphs:
        data_dict={}
        content=[]
        if val.text.strip()!="":
            # print(val.style.builtin)
            typ=val.style.name
            data_dict["type"]=typ


            data_dict["content"]=[]
            content=[]
            dic={}
            # for x in :
            # print(val.paragraph_format.alignment)

            for i in val.runs:
                # mapping font Style
                dic['style']=checkFontStyles(i.font)
                dic['text']=i.text
                # print(i.text.alignment)
                # adding font size
                if(i.font.size is not None):
                    dic['style']['fontSize']=float(i.font.size.pt)
                else:
                    dic['style']['fontSize']=float(0.0)
                content.append(dic.copy())


            # Merge duplicate Styles and compain Text
            data=mergeTextbyStyles(content)
            data_dict["content"]= data #append data to dictonary
            data_list.append(data_dict) #append dict to list
        
        #gathering image data
        if 'graphicData' in val._p.xml:
            dat=val._p.xml
            image=re.findall('image[0-9]*.[a-z]*',dat)
            # print(dat)
            if image!=[]:
                data_dict["type"]='Image'
                data_dict["content"]=image[0]
                # print(image_files)
                if image[0] not in image_files:
                    print('[*] Warning! '+image[0]+' not exists in destination folder')
                data_list.append(data_dict)
            else:
                print('[*] Unable to find Images ...')
    json_data=js.dumps(data_list,  indent=4)
    if store:
        f=open(json,'w')
        f.write(json_data)
        f.close
        print('[*] File successfully created')
        pwd=os.getcwd()
        print('[*] Saved To '+str(pwd)+' as '+str(json))
    else:
        print(json_data)
    return json_data


#-------------------------------------------
#          Manual Page
#-------------------------------------------


# word_to_json_parser() is the function for generating json file
# this function has 5 parameters:-
    #doc_path - must required, path of word file to parse
    #image_folder - optional, folder path to store the images of word file. default_value= docimages
    #store - optional, if given true the data will be stored in json file, if given false data will printed, but not stored
            #default=true
    #json - optional, json file name to store the generated data, default=output.json

#-------------------------------------------
#          Usage
#------------------------------------------- 
#   ./convert.py -/path/to/file.docx -o /path/to/outFile.json 


len_arg=len(sys.argv)
if(len_arg>=2 or len_arg ==3):
    filename= sys.argv[1]
    folderName=''.join(random.choice(letters) for i in range(10))
    outname=''
    store=False
    if(len_arg ==3):
        outname = sys.argv[2]
        folderName = sys.argv[2].split('.')[0]
        store=True
    word_to_json_parser(filename,store,folderName,outname)

else:
    print('Usage :\n')
    print('\t./{} inputFile.docx outputFile.json'.format(sys.argv[0].split('/')[-1]))
    print('\t\tOR')
    print('\t./{} inputFile.docx'.format(sys.argv[0].split('/')[-1]))





# ----------------------------------------
#       ToDo
# -----------------------------------------
#   1) Add Alignment (https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.paragraph.Paragraph.paragraph_format)
#   2) Add CmdLine Args
#
#




#-------------------------------------------
#           For Testing
#-------------------------------------------

# doc=docx.Document('MOD4.docx')
# # f=open('data.xml','w')
# j=0
# lst1=[]
# print(doc.author)
# for para in doc.paragraphs:
#     lst=[]
#     dic={}

#     # print(para.title)


#     for i in para.runs:

#         dic['style']=checkFontStyles(i.font)
#         dic['text']=i.text

#         # checkFontStyles(i.font)
#         if(i.font.size is not None):
#             dic['style']['fontSize']=float(i.font.size.pt)
#         else:
#             dic['style']['fontSize']=float(0.0)
        
        

#         lst.append(dic.copy())
#         # print(dic)
#     lst1.append(lst)
# for i in lst1:
#     mergeTextbyStyles(i)
